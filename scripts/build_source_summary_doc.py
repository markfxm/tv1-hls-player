from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


OUT = r"E:\tv\TV1\CCTV高清频道源测试总结_含超时表格.docx"


CHANNEL_SOURCES = [
    ("CCTV1", "节点1", "http://play.kankanlive.com/live/1698234869325962.m3u8"),
    ("CCTV1", "节点2", "http://58.57.40.22:9901/tsfile/live/0001_1.m3u8"),
    ("CCTV1", "节点3", "http://120.211.62.180:8000/hls/1/index.m3u8"),
    ("CCTV1", "节点4", "http://play.kankanlive.com/live/1661761962676984.m3u8"),
    ("CCTV2", "节点1", "https://play.kankanlive.com/live/1698234898628961.m3u8"),
    ("CCTV2", "节点2", "http://58.57.40.22:9901/tsfile/live/1001_1.m3u8"),
    ("CCTV2", "节点3", "http://120.211.62.180:8000/hls/2/index.m3u8"),
    ("CCTV2", "节点4", "http://183.11.239.36:808/hls/20/index.m3u8"),
    ("CCTV3", "节点1", "http://gmxw.7766.org:808/hls/91/index.m3u8"),
    ("CCTV3", "节点2", "http://183.11.239.36:808/hls/91/index.m3u8"),
    ("CCTV3", "节点3", "http://112.46.85.60:8009/hls/3/index.m3u8"),
    ("CCTV3", "节点4", "https://piccpndali.v.myalicdn.com/audio/cctv3_2.m3u8"),
    ("CCTV4", "节点1", "https://play.kankanlive.com/live/1698234936258960.m3u8"),
    ("CCTV4", "节点2", "http://58.57.40.22:9901/tsfile/live/1003_1.m3u8"),
    ("CCTV4", "节点3", "http://120.211.62.180:8000/hls/4/index.m3u8"),
    ("CCTV4", "节点4", "http://183.11.239.36:808/hls/22/index.m3u8"),
    ("CCTV5", "节点1", "http://gmxw.7766.org:808/hls/93/index.m3u8"),
    ("CCTV5", "节点2", "http://183.11.239.36:808/hls/93/index.m3u8"),
    ("CCTV5", "节点3", "http://112.46.85.60:8009/hls/503/index.m3u8"),
    ("CCTV5", "节点4", "https://piccpndali.v.myalicdn.com/audio/cctv5_2.m3u8"),
    ("CCTV6", "节点1", "http://58.57.40.22:9901/tsfile/live/1006_1.m3u8"),
    ("CCTV6", "节点2", "http://183.11.239.36:808/hls/94/index.m3u8"),
    ("CCTV6", "节点3", "http://182.150.23.74:808/hls/6/index.m3u8"),
    ("CCTV6", "节点4", "https://piccpndali.v.myalicdn.com/audio/cctv6_2.m3u8"),
    ("CCTV7", "节点1", "http://58.57.40.22:9901/tsfile/live/1007_1.m3u8"),
    ("CCTV7", "节点2", "http://120.211.62.180:8000/hls/7/index.m3u8"),
    ("CCTV7", "节点3", "http://8.138.7.223/tv/cctv7.m3u8"),
    ("CCTV7", "节点4", "http://183.11.239.36:808/hls/25/index.m3u8"),
    ("CCTV8", "节点1", "http://58.57.40.22:9901/tsfile/live/1008_1.m3u8"),
    ("CCTV8", "节点2", "http://gmxw.7766.org:808/hls/96/index.m3u8"),
    ("CCTV8", "节点3", "http://183.11.239.36:808/hls/96/index.m3u8"),
    ("CCTV8", "节点4", "http://112.46.85.60:8009/hls/8/index.m3u8"),
    ("CCTV9", "节点1", "https://play.kankanlive.com/live/1698423397390920.m3u8"),
    ("CCTV9", "节点2", "http://58.57.40.22:9901/tsfile/live/1009_1.m3u8"),
    ("CCTV9", "节点3", "http://120.211.62.180:8000/hls/9/index.m3u8"),
    ("CCTV9", "节点4", "http://183.11.239.36:808/hls/27/index.m3u8"),
    ("CCTV10", "节点1", "https://play.kankanlive.com/live/1698423445959919.m3u8"),
    ("CCTV10", "节点2", "http://120.211.62.180:8000/hls/10/index.m3u8"),
    ("CCTV10", "节点3", "http://183.11.239.36:808/hls/28/index.m3u8"),
    ("CCTV10", "节点4", "https://piccpndali.v.myalicdn.com/audio/cctv10_2.m3u8"),
    ("CCTV11", "节点1", "https://play.kankanlive.com/live/1698423476198918.m3u8"),
    ("CCTV11", "节点2", "http://58.57.40.22:9901/tsfile/live/1011_1.m3u8"),
    ("CCTV11", "节点3", "http://120.211.62.180:8000/hls/11/index.m3u8"),
    ("CCTV11", "节点4", "http://183.11.239.36:808/hls/29/index.m3u8"),
    ("CCTV12", "节点1", "https://play.kankanlive.com/live/1698423511884917.m3u8"),
    ("CCTV12", "节点2", "http://58.57.40.22:9901/tsfile/live/1012_1.m3u8"),
    ("CCTV12", "节点3", "http://120.211.62.180:8000/hls/12/index.m3u8"),
    ("CCTV12", "节点4", "http://183.11.239.36:808/hls/30/index.m3u8"),
    ("CCTV13", "节点1", "https://play.kankanlive.com/live/1698423543275916.m3u8"),
    ("CCTV13", "节点2", "https://event.pull.hebtv.com/jishi/cp1.m3u8"),
    ("CCTV13", "节点3", "http://58.57.40.22:9901/tsfile/live/1013_1.m3u8"),
    ("CCTV13", "节点4", "http://120.211.62.180:8000/hls/13/index.m3u8"),
    ("CCTV14", "节点1", "https://play.kankanlive.com/live/1698423575756915.m3u8"),
    ("CCTV14", "节点2", "https://event.pull.hebtv.com/jishi/cp2.m3u8"),
    ("CCTV14", "节点3", "http://58.57.40.22:9901/tsfile/live/1014_1.m3u8"),
    ("CCTV14", "节点4", "http://120.211.62.180:8000/hls/14/index.m3u8"),
    ("CCTV15", "节点1", "https://play.kankanlive.com/live/1698423607826914.m3u8"),
    ("CCTV15", "节点2", "http://58.57.40.22:9901/tsfile/live/1015_1.m3u8"),
    ("CCTV15", "节点3", "http://120.211.62.180:8000/hls/15/index.m3u8"),
    ("CCTV15", "节点4", "http://gmxw.7766.org:808/hls/102/index.m3u8"),
    ("CCTV16", "节点1", "http://gmxw.7766.org:808/hls/169/index.m3u8"),
    ("CCTV16", "节点2", "http://183.11.239.36:808/hls/169/index.m3u8"),
    ("CCTV16", "节点3", "https://piccpndali.v.myalicdn.com/audio/cctv16_2.m3u8"),
    ("CCTV17", "节点1", "https://play.kankanlive.com/live/1698423272597921.m3u8"),
    ("CCTV17", "节点2", "http://58.57.40.22:9901/tsfile/live/0019_1.m3u8"),
    ("CCTV17", "节点3", "http://120.211.62.180:8000/hls/17/index.m3u8"),
    ("CCTV17", "节点4", "http://hmfs.f3322.net:3388/hls/37/index.m3u8"),
]

HD_CANDIDATES = [
    ("CCTV4", "GITV HD 候选", "http://39.130.202.81:6610/gitv_live/G_CCTV-4-HD/G_CCTV-4-HD.m3u8", "超时"),
    ("CCTV7", "GITV HD 候选", "http://39.130.202.81:6610/gitv_live/G_CCTV-7-HD/G_CCTV-7-HD.m3u8", "超时"),
    ("CCTV9", "GITV HD 候选", "http://39.130.202.81:6610/gitv_live/G_CCTV-9-HD/G_CCTV-9-HD.m3u8", "超时"),
    ("CCTV10", "GITV HD 候选", "http://39.130.202.81:6610/gitv_live/G_CCTV-10-HD/G_CCTV-10-HD.m3u8", "超时"),
    ("CCTV12", "GITV HD 候选", "http://39.130.202.81:6610/gitv_live/G_CCTV-12-HD/G_CCTV-12-HD.m3u8", "超时"),
    ("CCTV14", "GITV HD 候选", "http://39.130.202.81:6610/gitv_live/G_CCTV-14-HD/G_CCTV-14-HD.m3u8", "超时"),
    ("CCTV16", "央视 liveop HD 候选", "http://liveop.cctv.cn/hls/CCTV16HD/playlist.m3u8", "超时"),
    ("CCTV17", "GITV HD 候选", "http://39.130.202.81:6610/gitv_live/G_CCTV-17-HD/G_CCTV-17-HD.m3u8", "超时"),
]


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text, bold=False, color=None):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(9)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, header in enumerate(headers):
        set_cell_text(hdr[i], header, bold=True)
        set_cell_shading(hdr[i], "E8EEF5")
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value)
    doc.add_paragraph()
    return table


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(10.5)


def add_source_line(doc, channel, label, url, status=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    prefix = p.add_run(f"{channel}：")
    prefix.bold = True
    prefix.font.name = "Microsoft YaHei"
    prefix._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    prefix.font.size = Pt(10)

    text = f"{label}({url})"
    if status:
        text += f" - {status}"
    run = p.add_run(text)
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(9)


def main():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)

    for style_name, size, color in [
        ("Heading 1", 16, "2E74B5"),
        ("Heading 2", 13, "2E74B5"),
        ("Heading 3", 12, "1F4D78"),
    ]:
        style = styles[style_name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(12)
        style.paragraph_format.space_after = Pt(6)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title.add_run("CCTV 高清频道源测试总结")
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(22)
    run.bold = True
    run.font.color.rgb = RGBColor.from_string("0B2545")

    meta = doc.add_paragraph()
    meta.add_run("生成位置：E:\\tv\\TV1    用途：为网页播放器筛选更清晰的 CCTV 频道节点").italic = True

    doc.add_heading("结论摘要", level=1)
    add_bullet(doc, "本轮搜索发现了若干标注 HD/高清的公开 CCTV 源，但多个高码率候选在当前网络环境下超时，暂不适合直接替换到播放器。")
    add_bullet(doc, "fanmingming 的 IPv6 频道列表可以正常下载，适合作为继续筛选的源列表，但其中多数源需要进一步逐个试播确认清晰度和兼容性。")
    add_bullet(doc, "当前播放器应优先采用实际能在浏览器播放的节点；标注 HD 不等于一定可用，也不等于 Chrome 能解码。")

    doc.add_heading("当前频道列表节点", level=1)
    doc.add_paragraph("以下按用户要求整理为“CCTVx：节点说明(链接)”格式。")
    for channel, label, url in CHANNEL_SOURCES:
        add_source_line(doc, channel, label, url)

    doc.add_heading("高清候选源测试记录", level=1)
    for channel, label, url, status in HD_CANDIDATES:
        add_source_line(doc, channel, label, url, f"测试结果：{status}")

    doc.add_heading("本机测试超时链接源表格", level=1)
    add_table(
        doc,
        ["频道", "来源", "链接", "本机测试结果", "建议"],
        [
            [channel, label, url, status, "当前网络环境下暂不采用，后续可换网络重测"]
            for channel, label, url, status in HD_CANDIDATES
        ],
    )

    doc.add_heading("已找到的列表来源", level=1)
    add_table(
        doc,
        ["列表名称", "地址", "状态", "备注"],
        [
            [
                "fanmingming IPv6 直播列表",
                "https://live.fanmingming.com/tv/m3u/ipv6.m3u",
                "可下载",
                "包含 CCTV-1 综合、CCTV-10 科教、CCTV-13 新闻、CCTV-16 奥林匹克等条目。",
            ],
            [
                "fanmingming GitHub IPv6 直播列表",
                "https://raw.githubusercontent.com/fanmingming/live/main/tv/m3u/ipv6.m3u",
                "可下载",
                "适合作为公网可访问备用入口。",
            ],
            [
                "TVBox live.txt",
                "https://raw.githubusercontent.com/yoursmile66/TVBox/main/live.txt",
                "此前可下载",
                "当前网页播放器里的 CCTV1-17 节点主要来自该公开列表。",
            ],
        ],
    )

    doc.add_heading("当前建议", level=1)
    add_bullet(doc, "不要一次性把所有“HD”候选替换进网页。先做批量试播脚本，记录 readyState、videoWidth、videoHeight、是否只有声音、是否超时。")
    add_bullet(doc, "优先保留 HTTPS 节点；HTTP 节点在浏览器环境更容易受网络、跨域或混合内容策略影响。")
    add_bullet(doc, "标注为 audio 的地址很可能只有声音，不建议作为优先节点。")
    add_bullet(doc, "如果后续要追求清晰度，可以把每个频道的节点按实测分辨率排序，而不是按来源文字排序。")

    doc.add_heading("下一步可执行动作", level=1)
    add_table(
        doc,
        ["动作", "目的", "输出"],
        [
            ["批量抓取 fanmingming IPv6 列表", "获得更多 CCTV 候选源", "按频道聚合的候选 URL"],
            ["自动试播每个候选源 8-10 秒", "过滤失效和只有声音的源", "readyState、分辨率、错误原因"],
            ["把高分辨率且可播放的源排到节点1", "改善默认播放清晰度", "更新 public/channels.json"],
            ["保留旧节点作为备用", "避免高清源不稳定时无源可看", "每个频道多个节点"],
        ],
    )

    doc.add_heading("注意事项", level=1)
    p = doc.add_paragraph()
    p.add_run("这些地址来自公开互联网列表，稳定性、版权授权、地区可访问性和长期可用性都无法保证。").bold = True
    doc.add_paragraph("实际用于产品或 APK 前，需要确认播放源的合法授权和可持续可访问性。")

    doc.save(OUT)
    print("created_docx")


if __name__ == "__main__":
    main()
